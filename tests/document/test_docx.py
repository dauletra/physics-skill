"""The Word backend: the package, the flat body, and what must not leak.

Two kinds of check, and both are needed. The structural ones read the XML we
wrote and are precise about the document's own promises — numbering, the
answers section, escaping. The independent one opens the file with a library
that was not written here: Word says nothing useful when it refuses a
package, so a second reader is the closest thing to Word available in CI.

Neither replaces opening the file by hand — see evals/word.md.
"""

from __future__ import annotations

import io
import os
import re
import zipfile
from pathlib import Path
from xml.dom import minidom
from xml.etree import ElementTree

import docx
import pytest

from kinds import EACH_KIND
from physics_svg.document import build_docx, parse_block, parse_document
from physics_svg.document.emit.docx.wml import P_ORDER, R_ORDER, props

REPO = Path(__file__).resolve().parent.parent.parent
GOLDEN = REPO / "tests" / "golden" / "document"
REGENERATE = os.environ.get("REGEN_GOLDEN") == "1"

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

DOC = parse_document({"title": "Проба"})
TEXT = {"type": "text", "body": "Условие."}
OPEN = {"type": "open", "answer": "42"}


def task(*blocks: dict, **extra: object) -> dict:
    return {"type": "task", "blocks": list(blocks), **extra}


def build(*raw_blocks: dict, with_answers: bool = True) -> bytes:
    blocks = [parse_block(raw, f"b{i}") for i, raw in enumerate(raw_blocks)]
    return build_docx(DOC, blocks, with_answers=with_answers)


def body(data: bytes) -> ElementTree.Element:
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    found = root.find(f"{W}body")
    assert found is not None
    return found


def texts(data: bytes) -> list[str]:
    """The text of every paragraph, tables included — `w:t` in document order,
    grouped by the paragraph it belongs to."""
    return ["".join(node.itertext()) for node in body(data).iter(f"{W}p")]


def _numbers(lines: list[str]) -> list[str]:
    """The task numbers actually printed on the sheet."""
    return [match.group() for line in lines if (match := re.match(r"^\d+\.", line))]


class TestPackage:
    def test_it_is_a_zip_with_the_parts_word_needs(self) -> None:
        with zipfile.ZipFile(io.BytesIO(build(task(TEXT, OPEN)))) as package:
            names = set(package.namelist())
        assert {
            "[Content_Types].xml",
            "_rels/.rels",
            "word/_rels/document.xml.rels",
            "word/document.xml",
            "word/styles.xml",
            "word/settings.xml",
            "docProps/core.xml",
        } <= names

    def test_every_part_is_declared_in_the_content_types(self) -> None:
        # A part Word cannot type is a part Word refuses to read.
        with zipfile.ZipFile(io.BytesIO(build(task(TEXT, OPEN)))) as package:
            types = package.read("[Content_Types].xml").decode()
            parts = [name for name in package.namelist() if name.endswith(".xml")]
        for part in parts:
            if part == "[Content_Types].xml":
                continue
            assert f'PartName="/{part}"' in types, part

    def test_every_relationship_points_at_a_part(self) -> None:
        with zipfile.ZipFile(io.BytesIO(build(task(TEXT, OPEN)))) as package:
            names = set(package.namelist())
            for rels in ("_rels/.rels", "word/_rels/document.xml.rels"):
                base = rels.rsplit("_rels/", 1)[0]
                for target in re.findall(r'Target="([^"]+)"', package.read(rels).decode()):
                    assert f"{base}{target}" in names, target

    def test_two_builds_are_byte_identical(self) -> None:
        # Without this a golden of a binary format is impossible, and so is
        # telling a real change from a timestamp.
        assert build(task(TEXT, OPEN)) == build(task(TEXT, OPEN))

    def test_the_title_reaches_the_file_properties(self) -> None:
        # And not the sheet: what the sheet is called is a `heading` block.
        with zipfile.ZipFile(io.BytesIO(build(task(TEXT, OPEN)))) as package:
            assert "<dc:title>Проба</dc:title>" in package.read("docProps/core.xml").decode()
        assert "Проба" not in "".join(texts(build(task(TEXT, OPEN))))

    def test_the_section_closes_the_body(self) -> None:
        # `w:sectPr` anywhere but last is a schema violation.
        children = list(body(build(task(TEXT, OPEN))))
        assert children[-1].tag == f"{W}sectPr"


class TestIndependentReader:
    """A reader that was not written here. Word gives no diagnostics at all
    when it rejects a package; this is the closest CI can get to it."""

    def test_the_document_opens(self) -> None:
        document = docx.Document(io.BytesIO(build(task(TEXT, OPEN))))
        assert [p.text for p in document.paragraphs]

    def test_named_styles_are_the_ones_the_teacher_sees(self) -> None:
        data = build({"type": "heading", "text": "Кинематика", "level": 1}, task(TEXT, OPEN))
        document = docx.Document(io.BytesIO(data))
        used = {p.style.name for p in document.paragraphs}
        assert "Заголовок листа" in used
        assert "Задание" in used

    def test_a_table_survives_the_round_trip(self) -> None:
        data = build(
            {"type": "table", "headers": ["t, с", "υ, м/с"], "rows": [["0", "5"], ["2", "9"]]}
        )
        table = docx.Document(io.BytesIO(data)).tables[0]
        assert [c.text for c in table.rows[0].cells] == ["t, с", "υ, м/с"]
        assert len(table.columns) == 2


class TestFlatBody:
    """What HTML says by nesting, Word has to say with paragraph properties."""

    def test_tasks_are_numbered_in_order(self) -> None:
        # Without the section, so that a caption in it cannot be mistaken for
        # a number on the sheet.
        lines = texts(build(task(TEXT, OPEN), task(TEXT, OPEN), with_answers=False))
        assert _numbers(lines) == ["1.", "2."]

    def test_theory_between_tasks_does_not_shift_the_numbering(self) -> None:
        lines = texts(
            build(
                task(TEXT, OPEN),
                {"type": "text", "body": "Теория."},
                task(TEXT, OPEN),
                with_answers=False,
            )
        )
        assert _numbers(lines) == ["1.", "2."]

    def test_a_subtask_is_lettered_and_indented_further_than_its_task(self) -> None:
        data = build(
            task(
                TEXT,
                {"type": "part", "blocks": [{"type": "text", "body": "Первое"}, OPEN]},
                {"type": "part", "blocks": [{"type": "text", "body": "Второе"}, OPEN]},
            )
        )
        indents = {}
        for node in body(data).iter(f"{W}p"):
            line = "".join(node.itertext())
            ind = node.find(f"{W}pPr/{W}ind")
            if ind is not None:
                indents[line] = int(ind.get(f"{W}left", "0"))
        assert any(line.startswith("а)") for line in indents)
        assert any(line.startswith("б)") for line in indents)
        task_line = next(value for line, value in indents.items() if line.startswith("1."))
        part_line = next(value for line, value in indents.items() if line.startswith("а)"))
        assert part_line > task_line

    def test_a_row_becomes_a_borderless_table(self) -> None:
        data = build({"type": "row", "blocks": [TEXT, {"type": "text", "body": "Рядом."}]})
        tables = list(body(data).iter(f"{W}tbl"))
        assert len(tables) == 1
        assert tables[0].find(f"{W}tblPr/{W}tblBorders") is None

    def test_a_wrapped_row_fills_columns_downwards(self) -> None:
        # The same promise the CSS grid makes: with four items in two columns,
        # the first two are on the left.
        items = [{"type": "text", "body": name} for name in ("раз", "два", "три", "четыре")]
        data = build({"type": "row", "columns": 2, "blocks": items})
        table = next(body(data).iter(f"{W}tbl"))
        rows = [
            ["".join(cell.itertext()) for cell in row.iter(f"{W}tc")]
            for row in table.iter(f"{W}tr")
        ]
        assert rows == [["раз", "три"], ["два", "четыре"]]

    def test_a_table_repeats_its_header_over_a_page_break(self) -> None:
        # Print is engineered here, unlike on screen — that is the point.
        data = build({"type": "table", "headers": ["a"], "rows": [["1"]]})
        first = next(body(data).iter(f"{W}tr"))
        assert first.find(f"{W}trPr/{W}tblHeader") is not None


class TestAnswers:
    def test_the_answer_reaches_the_section_and_not_the_body(self) -> None:
        full = "".join(texts(build(task(TEXT, OPEN))))
        assert "Ответы" in full and "42" in full
        handout = "".join(texts(build(task(TEXT, OPEN), with_answers=False)))
        assert "42" not in handout and "Ответы" not in handout

    def test_a_document_without_answers_prints_no_section(self) -> None:
        lines = "".join(texts(build({"type": "text", "body": "Конспект."})))
        assert "Ответы" not in lines

    def test_an_explanation_follows_its_answer(self) -> None:
        data = build(task(TEXT, {**OPEN, "explanation": "По определению."}))
        assert any("По определению." in line for line in texts(data))


class TestEscaping:
    def test_author_markup_never_becomes_xml(self) -> None:
        payload = '<w:p>&<>"</w:p>'
        data = build(task({"type": "text", "body": payload}, OPEN))
        with zipfile.ZipFile(io.BytesIO(data)) as package:
            document = package.read("word/document.xml").decode()
        assert "&lt;w:p&gt;" in document
        assert payload in "".join(texts(data))

    def test_control_characters_are_dropped(self) -> None:
        # XML 1.0 has no way to write them, and one of them makes the whole
        # package unopenable — a document that silently loses a byte is
        # better than a file the teacher cannot open.
        data = build(task({"type": "text", "body": "до\x07после"}, OPEN))
        assert any("допосле" in line for line in texts(data))
        docx.Document(io.BytesIO(data))

    def test_a_space_before_a_blank_survives(self) -> None:
        data = build({"type": "text", "body": "Дата: ______"})
        assert any(line.startswith("Дата: ") for line in texts(data))


class TestProperties:
    """`w:pPr` and `w:rPr` are sequences, not sets: an element in the wrong
    place makes Word reject the document without saying which one."""

    def test_properties_come_out_in_schema_order(self) -> None:
        given = {"w:ind": "<w:ind/>", "w:pStyle": "<w:pStyle/>", "w:spacing": "<w:spacing/>"}
        assert props("w:pPr", P_ORDER, given) == (
            "<w:pPr><w:pStyle/><w:spacing/><w:ind/></w:pPr>"
        )

    def test_underline_comes_after_size_and_vertalign_last(self) -> None:
        given = {"w:vertAlign": "<w:vertAlign/>", "w:u": "<w:u/>", "w:sz": "<w:sz/>"}
        assert props("w:rPr", R_ORDER, given) == "<w:rPr><w:sz/><w:u/><w:vertAlign/></w:rPr>"

    def test_an_unknown_property_is_refused(self) -> None:
        with pytest.raises(ValueError, match="unknown properties"):
            props("w:pPr", P_ORDER, {"w:outlineLvl": "<w:outlineLvl/>"})


class TestPictures:
    """An illustration is a group of native shapes, not a picture: the teacher
    can move a label and print it at any size."""

    def test_a_ruling_becomes_shapes(self) -> None:
        data = build({"type": "paper", "ruling": "grid", "rows": 3, "cols": 4})
        with zipfile.ZipFile(io.BytesIO(data)) as package:
            document = package.read("word/document.xml").decode()
        assert "<w:drawing>" in document
        assert "<wpg:wgp>" in document and document.count("<wps:wsp>") > 3

    def test_every_drawing_is_numbered_once(self) -> None:
        # Two drawings sharing a `docPr` id make Word refuse the document.
        data = build(
            {"type": "paper", "ruling": "lines", "rows": 2},
            {"type": "paper", "ruling": "lines", "rows": 3},
        )
        with zipfile.ZipFile(io.BytesIO(data)) as package:
            document = package.read("word/document.xml").decode()
        ids = re.findall(r'<wp:docPr id="(\d+)"', document)
        assert ids == sorted(ids, key=int) and len(set(ids)) == len(ids)

    def test_a_drawing_never_outgrows_its_column(self) -> None:
        # Inside a two-column row a picture has half the width to live in.
        wide = _extent(build({"type": "paper", "ruling": "lines", "rows": 2}))
        narrow = _extent(
            build(
                {
                    "type": "row",
                    "blocks": [
                        {"type": "paper", "ruling": "lines", "rows": 2},
                        TEXT,
                    ],
                }
            )
        )
        assert narrow[0] < wide[0]

    def test_a_plotted_answer_reaches_the_answers_section(self) -> None:
        data = build(
            task(
                TEXT,
                {
                    "type": "plot",
                    "x_label": "t, с",
                    "y_label": "υ, м/с",
                    "x_range": [0, 10],
                    "y_range": [0, 20],
                    "x_step": 2,
                    "y_step": 5,
                    "answer": [{"points": [[0, 0], [10, 20]]}],
                },
            )
        )
        with zipfile.ZipFile(io.BytesIO(data)) as package:
            document = package.read("word/document.xml").decode()
        # Two drawings: the empty axes in the task, the answer at the end.
        assert document.count("<w:drawing>") == 2
        assert "Ответы" in "".join(texts(data))


def _extent(data: bytes) -> tuple[int, int]:
    with zipfile.ZipFile(io.BytesIO(data)) as package:
        document = package.read("word/document.xml").decode()
    match = re.search(r'<wp:extent cx="(\d+)" cy="(\d+)"/>', document)
    assert match is not None
    return int(match.group(1)), int(match.group(2))


@EACH_KIND
class TestEveryKind:
    def test_the_documented_example_builds(self, kind) -> None:
        """Every kind prints — no exceptions and no allowances.

        Until illustrations arrived this test also accepted a refusal; it does
        not any more, which is what keeps a kind from quietly working in one
        format only.
        """
        for i, example in enumerate(kind.tasks):
            data = build_docx(DOC, [parse_block(example, f"{kind.tag}-{i}")])
            docx.Document(io.BytesIO(data))
            assert "".join(texts(data)).strip()

    def test_the_answer_does_not_leak_into_the_body(self, kind) -> None:
        for i, example in enumerate(kind.tasks):
            block = parse_block(example, f"{kind.tag}-{i}")
            data = build_docx(DOC, [block], with_answers=False)
            assert "Ответы" not in "".join(texts(data))


class TestGolden:
    def test_a_mixed_document_matches_its_golden(self) -> None:
        data = build(
            {"type": "heading", "text": "Кинематика", "level": 1},
            {"type": "divider"},
            task(TEXT, OPEN, {"type": "answer_line"}),
            {"type": "row", "blocks": [TEXT, {"type": "list", "items": ["раз", "два"]}]},
            task(
                {"type": "text", "body": "Скорость υ<sub>0</sub> и дата: ______"},
                {"type": "part", "blocks": [{"type": "text", "body": "а"}, OPEN]},
                {
                    "type": "part",
                    "blocks": [
                        {"type": "text", "body": "б"},
                        {
                            "type": "choice",
                            "options": [{"text": "верно", "correct": True}, {"text": "нет"}],
                        },
                    ],
                },
            ),
            {"type": "table", "headers": ["t, с", "υ, м/с"], "rows": [["0", "5"]]},
        )
        with zipfile.ZipFile(io.BytesIO(data)) as package:
            document = package.read("word/document.xml").decode()
        # The golden is the formatted XML, not the zip: a binary golden cannot
        # be read in a diff, and the whole point of freezing it is being able
        # to look at what changed.
        pretty = minidom.parseString(document).toprettyxml(indent="  ")
        path = GOLDEN / "mixed.docx.xml"
        if REGENERATE or not path.exists():
            path.write_text(pretty, encoding="utf-8")
            if not REGENERATE:
                return
        assert pretty == path.read_text(encoding="utf-8"), (
            "документ Word разошёлся с эталоном mixed.docx.xml; если изменение "
            "намеренное — пересними эталоны и просмотри диф глазами"
        )
