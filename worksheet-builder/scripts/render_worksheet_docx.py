#!/usr/bin/env python3
"""Render a worksheet draft (meta.json + tasks/*.json) into an editable Word
(.docx) document.

This is a best-effort approximation of the HTML/print worksheet produced by
render_worksheet.py, not a pixel copy — Word has no CSS flex/grid, no live
SVG, no KaTeX. See "Экспорт в Word" in references/design-system.md for the
deliberate simplifications (no interactive toolbar, charts/illustrations
become static raster images, LaTeX formulas become images too, chart legends
collapse to a plain text line).

Requires the packages in requirements-docx.txt (pip install -r
requirements-docx.txt) — python-docx, matplotlib, svglib. Not needed at all
for the plain HTML workflow; only imported when --docx is passed to
render_worksheet.py (or when this module is run directly, mainly for
development/testing).

Usage:
    python render_worksheet_docx.py <workspace_dir>

Normally invoked via `python render_worksheet.py <workspace_dir> --docx`.
"""
import io
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt
from docx.table import _Cell

import matplotlib

matplotlib.use("Agg")
from matplotlib import mathtext  # noqa: E402
from matplotlib.font_manager import FontProperties  # noqa: E402

from PIL import Image  # noqa: E402 (pulled in transitively by matplotlib)
from svglib.svglib import svg2rlg  # noqa: E402
from reportlab.graphics import renderPM  # noqa: E402

from render_worksheet import (  # noqa: E402
    LETTERS,
    LOWER_LETTERS,
    SVG_SNIPPETS,
    build_chart_svg,
    load_workspace,
    t,
)

PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
MARGIN_TOP_MM = 15
MARGIN_BOTTOM_MM = 15
MARGIN_RIGHT_MM = 15
MARGIN_LEFT_MM = 18  # extra room for binding, mirrors BASE_CSS's @page margin
CONTENT_WIDTH_MM = PAGE_WIDTH_MM - MARGIN_LEFT_MM - MARGIN_RIGHT_MM

FONT_NAME = "PT Sans"
HEADER_INDENT_MM = 8

SIDE_VISUAL_TYPES = {"chart", "illustration"}

_LEGEND_DASH = {"solid": "———", "dashed": "– – –", "dotted": "· · ·"}
_MATH_RE = re.compile(r"\$\$(.+?)\$\$|\$(.+?)\$", re.DOTALL)
_SUPSUB_RE = re.compile(r"<(sup|sub)>(.*?)</\1>", re.DOTALL)
_SVG_TAG_RE = re.compile(r"(<svg.*?</svg>)", re.DOTALL)
_FILL_BLANK_RE = re.compile(r"___(\w+)___")


# ---------------------------------------------------------------------------
# Low-level OOXML helpers python-docx doesn't expose directly.
# ---------------------------------------------------------------------------

def _set_paragraph_border(paragraph, sides, val="single", size=6, space=1, color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), val)
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), color)
        pBdr.append(el)


def _set_paragraph_shading(paragraph, fill):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def _set_hanging(paragraph, indent_mm=HEADER_INDENT_MM):
    fmt = paragraph.paragraph_format
    fmt.left_indent = Mm(indent_mm)
    fmt.first_line_indent = Mm(-indent_mm)
    fmt.tab_stops.add_tab_stop(Mm(indent_mm), WD_TAB_ALIGNMENT.LEFT)


# ---------------------------------------------------------------------------
# Rich text: text -> <sup>/<sub> -> LaTeX, same priority order as symbols.md
# and esc()/SUPSUB_TAGS in render_worksheet.py, just targeting docx runs
# instead of an HTML string.
# ---------------------------------------------------------------------------

def _split_math(text):
    tokens = []
    pos = 0
    for m in _MATH_RE.finditer(text):
        if m.start() > pos:
            tokens.append(("text", text[pos:m.start()]))
        if m.group(1) is not None:
            tokens.append(("display_math", m.group(1)))
        else:
            tokens.append(("inline_math", m.group(2)))
        pos = m.end()
    if pos < len(text):
        tokens.append(("text", text[pos:]))
    return tokens


def _add_plain_run(paragraph, text, bold, italic, size):
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.bold = bold
    if italic:
        run.font.italic = True
    if size:
        run.font.size = size


def _add_text_with_supsub(paragraph, text, bold, italic, size):
    pos = 0
    for m in _SUPSUB_RE.finditer(text):
        if m.start() > pos:
            _add_plain_run(paragraph, text[pos:m.start()], bold, italic, size)
        run = paragraph.add_run(m.group(2))
        run.font.bold = bold
        if italic:
            run.font.italic = True
        if size:
            run.font.size = size
        if m.group(1) == "sup":
            run.font.superscript = True
        else:
            run.font.subscript = True
        pos = m.end()
    if pos < len(text):
        _add_plain_run(paragraph, text[pos:], bold, italic, size)


_MATH_IMAGE_CACHE = {}


def _math_png_bytes(latex, fontsize_pt):
    key = (latex, round(fontsize_pt, 1))
    if key not in _MATH_IMAGE_CACHE:
        buf = io.BytesIO()
        prop = FontProperties(size=fontsize_pt)
        mathtext.math_to_image(f"${latex}$", buf, prop=prop, dpi=300, format="png")
        buf.seek(0)
        _MATH_IMAGE_CACHE[key] = buf.read()
    return _MATH_IMAGE_CACHE[key]


def _add_math_image(paragraph, latex, size):
    fontsize_pt = size.pt if size else 12
    try:
        png_bytes = _math_png_bytes(latex.strip(), fontsize_pt)
    except Exception:
        # mathtext doesn't cover 100% of KaTeX/LaTeX syntax (e.g. some \text{}
        # edge cases) — degrade to raw "$...$" text instead of failing the
        # whole build over one formula.
        run = paragraph.add_run(f"${latex.strip()}$")
        if size:
            run.font.size = size
        return
    img = Image.open(io.BytesIO(png_bytes))
    width_px, height_px = img.size
    height_emu = Emu(int(fontsize_pt * 1.35 * 12700))
    width_emu = Emu(int(height_emu * width_px / height_px))
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(png_bytes), height=height_emu, width=width_emu)


def _add_rich_text(paragraph, text, *, bold=False, italic=False, size=None):
    if not text:
        return
    for kind, content in _split_math(str(text)):
        if kind == "text":
            _add_text_with_supsub(paragraph, content, bold, italic, size)
        else:
            _add_math_image(paragraph, content, size)


def _add_underline_blank(paragraph, value, min_chars, size=None):
    text = value or ""
    pad = max(min_chars - len(text), 2)
    run = paragraph.add_run(text + " " * pad)
    run.underline = True
    if size:
        run.font.size = size


# ---------------------------------------------------------------------------
# SVG (chart / illustration) -> raster image, via the exact same SVG-building
# code the HTML renderer uses (build_chart_svg / SVG_SNIPPETS), so the two
# outputs never drift apart.
# ---------------------------------------------------------------------------

def _extract_svg(wrapped_markup):
    m = _SVG_TAG_RE.search(wrapped_markup)
    if not m:
        raise ValueError("Expected an <svg> element in generated markup")
    return m.group(1)


def _svg_to_png_bytes(svg_str, scale=4):
    drawing = svg2rlg(io.StringIO(svg_str))
    if drawing is None:
        raise ValueError("svglib failed to parse generated SVG")
    drawing.width *= scale
    drawing.height *= scale
    drawing.scale(scale, scale)
    buf = io.BytesIO()
    renderPM.drawToFile(drawing, buf, fmt="PNG", bg=0xFFFFFF)
    buf.seek(0)
    return buf.read()


def _container_width_mm(container):
    if isinstance(container, _Cell) and container.width:
        return container.width.mm
    return CONTENT_WIDTH_MM


def _add_svg_image(container, svg_str):
    png_bytes = _svg_to_png_bytes(svg_str)
    img = Image.open(io.BytesIO(png_bytes))
    width_px, height_px = img.size
    width_mm = _container_width_mm(container) * 0.95
    width = Mm(width_mm)
    height = Emu(int(width * height_px / width_px))
    p = container.add_paragraph()
    run = p.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=width, height=height)


def _chart_legend_text(spec):
    parts = []
    for s in spec.get("series", []):
        if s.get("label"):
            dash = _LEGEND_DASH.get(s.get("style", "solid"), "———")
            parts.append(f'{dash} {s["label"]}')
    return "     ".join(parts)


# ---------------------------------------------------------------------------
# Per-task-type renderers. Each takes (container, task, is_teacher, lang) and
# returns the answer to show in the teacher version, or None — mirrors
# TASK_RENDERERS in render_worksheet.py. `container` is either the Document
# (full-width tasks) or a table _Cell (chart/illustration side layout) — both
# expose the same add_paragraph()/add_table() API, so renderers don't need to
# know which one they got.
# ---------------------------------------------------------------------------

def _render_text_with_solution(container, task, is_teacher, lang):
    lines = int(task.get("solution_lines", 4))
    if task.get("show_solution", True):
        for _ in range(lines):
            p = container.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.add_run(" ")
            _set_paragraph_border(p, ["bottom"])
    return task.get("answer")


def _render_text_no_solution(container, task, is_teacher, lang):
    if task.get("answer_blank", True):
        p = container.add_paragraph()
        p.add_run(f'{t(lang, "answer_label")} ')
        _add_underline_blank(p, "", min_chars=20)
    return task.get("answer")


def _render_matching(container, task, is_teacher, lang):
    left_items = task.get("left_items", [])
    right_items = task.get("right_items", [])
    answer_map = task.get("answer_map", {})
    n = max(len(left_items), len(right_items))
    table = container.add_table(rows=n, cols=2)
    for i in range(n):
        if i < len(left_items):
            _add_rich_text(table.cell(i, 0).paragraphs[0], f"{i + 1}. {left_items[i]}")
        if i < len(right_items):
            _add_rich_text(table.cell(i, 1).paragraphs[0], f"{LETTERS[i]}. {right_items[i]}")
    if not answer_map:
        return None
    letters = {item: LETTERS[i] for i, item in enumerate(right_items)}
    pairs = [f"{i}-{letters.get(answer_map.get(item), '?')}" for i, item in enumerate(left_items, start=1)]
    return ", ".join(pairs)


def _render_true_false(container, task, is_teacher, lang):
    for st in task.get("statements", []):
        p = container.add_paragraph()
        _add_rich_text(p, st.get("text", ""))
        correct = st.get("answer")
        true_mark = "☒" if is_teacher and correct is True else "☐"
        false_mark = "☒" if is_teacher and correct is False else "☐"
        p.add_run(f'   {true_mark} {t(lang, "true_label")}   {false_mark} {t(lang, "false_label")}')
    return None


def _render_multiple_choice(container, task, is_teacher, lang):
    options = task.get("options", [])
    correct = task.get("correct")
    for i, opt in enumerate(options):
        p = container.add_paragraph()
        is_correct = is_teacher and correct is not None and i == correct
        prefix = "✓ " if is_correct else "   "
        run = p.add_run(f"{prefix}{LOWER_LETTERS[i]}) ")
        run.font.bold = is_correct
        _add_rich_text(p, opt, bold=is_correct)
    return task.get("answer")


def _render_fill_blank_text(container, task, is_teacher, lang):
    template = task.get("text_template", "")
    blanks = task.get("blanks", {})
    p = container.add_paragraph()
    pos = 0
    for m in _FILL_BLANK_RE.finditer(template):
        _add_rich_text(p, template[pos:m.start()])
        key = m.group(1)
        if is_teacher and key in blanks:
            run = p.add_run(str(blanks[key]))
            run.font.bold = True
            run.underline = True
        else:
            run = p.add_run(" " * 8)
            run.underline = True
        pos = m.end()
    _add_rich_text(p, template[pos:])
    return None


def _render_fill_blank_table(container, task, is_teacher, lang):
    headers = task.get("headers", [])
    rows = task.get("rows", [])
    answers = task.get("answers", {})
    table = container.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    for c, h in enumerate(headers):
        run = table.cell(0, c).paragraphs[0].add_run(str(h))
        run.font.bold = True
    for r_i, row in enumerate(rows):
        for c_i, cell in enumerate(row):
            p = table.cell(r_i + 1, c_i).paragraphs[0]
            if cell is None:
                value = answers.get(f"{r_i},{c_i}", "") if is_teacher else ""
                if value:
                    run = p.add_run(str(value))
                    run.font.bold = True
            else:
                p.add_run(str(cell))
    return None


def _render_chart(container, task, is_teacher, lang):
    spec = task["chart_spec"]
    svg = _extract_svg(build_chart_svg(spec))
    _add_svg_image(container, svg)
    legend = _chart_legend_text(spec)
    if legend:
        run = container.add_paragraph().add_run(legend)
        run.font.size = Pt(8)
    return task.get("answer")


def _render_illustration(container, task, is_teacher, lang):
    if "svg_snippet" in task:
        name = task["svg_snippet"]["name"]
        params = task["svg_snippet"].get("params", {})
        if name not in SVG_SNIPPETS:
            raise ValueError(f"Unknown svg_snippet '{name}'. Available: {list(SVG_SNIPPETS)}.")
        svg = SVG_SNIPPETS[name](params)
    elif "raw_svg" in task:
        svg = task["raw_svg"]
    else:
        raise ValueError(f"Illustration task {task.get('id')} needs svg_snippet or raw_svg")
    _add_svg_image(container, svg)
    return task.get("answer")


TASK_RENDERERS_DOCX = {
    "text_with_solution": _render_text_with_solution,
    "text_no_solution": _render_text_no_solution,
    "matching": _render_matching,
    "true_false": _render_true_false,
    "fill_blank_text": _render_fill_blank_text,
    "fill_blank_table": _render_fill_blank_table,
    "chart": _render_chart,
    "illustration": _render_illustration,
    "multiple_choice": _render_multiple_choice,
}


# ---------------------------------------------------------------------------
# Task assembly: number/letter + prompt header, body, answer block. Shared by
# both top-level tasks and compound parts (indent_mm distinguishes them) so
# the chart/illustration side-layout logic exists in exactly one place.
# ---------------------------------------------------------------------------

def _render_header_line(paragraph, num_label, prompt, points, lang):
    _set_hanging(paragraph)
    run = paragraph.add_run(f"{num_label}\t")
    run.font.bold = True
    _add_rich_text(paragraph, prompt)
    if points:
        run = paragraph.add_run(f' ({points} {t(lang, "points_suffix")})')
        run.font.size = Pt(9)


def _render_answer_block(container, answer, lang, indent_mm=0):
    p = container.add_paragraph()
    p.paragraph_format.left_indent = Mm(indent_mm)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{t(lang, "answer_label")} ')
    run.font.bold = True
    _add_rich_text(p, str(answer))
    _set_paragraph_shading(p, fill="F2F2F2")
    _set_paragraph_border(p, ["top", "bottom", "left", "right"], val="dashed", space=4)


def _visual_column_widths_mm(visual_width_pct):
    pct = int(visual_width_pct)
    right = CONTENT_WIDTH_MM * pct / 100
    return [CONTENT_WIDTH_MM - right, right]


def _set_table_widths(table, widths_mm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            cell.width = Mm(w)
    for col, w in zip(table.columns, widths_mm):
        col.width = Mm(w)


def _render_task_body(document, ttype, renderer, num_label, prompt, points, task_data, is_teacher, lang, indent_mm=0):
    if ttype in SIDE_VISUAL_TYPES and task_data.get("visual_position", "side") != "below":
        # Text (left) / chart-or-illustration image (right), same ~36%-by-default
        # split as .task-with-visual in BASE_CSS. Not indented for compound
        # parts (see design-system.md "Экспорт в Word") — table-level indent
        # needs raw OOXML (w:tblInd) that isn't worth the complexity here.
        table = document.add_table(rows=1, cols=2)
        _set_table_widths(table, _visual_column_widths_mm(task_data.get("visual_width", "36")))
        left_cell, right_cell = table.rows[0].cells
        _render_header_line(left_cell.paragraphs[0], num_label, prompt, points, lang)
        answer = renderer(right_cell, task_data, is_teacher, lang)
        if is_teacher and answer:
            _render_answer_block(left_cell, answer, lang)
        return

    header_p = document.add_paragraph()
    header_p.paragraph_format.left_indent = Mm(indent_mm)
    _render_header_line(header_p, num_label, prompt, points, lang)
    answer = renderer(document, task_data, is_teacher, lang)
    if is_teacher and answer:
        _render_answer_block(document, answer, lang, indent_mm=indent_mm)


def render_task_docx(document, index, task, is_teacher, lang):
    ttype = task.get("type")
    if ttype == "compound":
        render_compound_task_docx(document, index, task, is_teacher, lang)
        return
    renderer = TASK_RENDERERS_DOCX.get(ttype)
    if renderer is None:
        raise ValueError(f"Unknown task type '{ttype}' in task {task.get('id')}")
    _render_task_body(
        document, ttype, renderer, f"{index}.", task.get("prompt", ""), task.get("points"), task, is_teacher, lang
    )


def render_compound_task_docx(document, index, task, is_teacher, lang):
    header_p = document.add_paragraph()
    _render_header_line(header_p, f"{index}.", task.get("prompt", ""), task.get("points"), lang)
    for i, part in enumerate(task.get("parts", [])):
        ptype = part.get("type")
        if ptype == "compound":
            raise ValueError(f"Nested 'compound' part not allowed in task {task.get('id')}")
        renderer = TASK_RENDERERS_DOCX.get(ptype)
        if renderer is None:
            raise ValueError(f"Unknown part type '{ptype}' in task {task.get('id')} (part {i})")
        label = part.get("label") or LOWER_LETTERS[i]
        _render_task_body(
            document, ptype, renderer, f"{label})", part.get("prompt", ""), part.get("points"),
            part, is_teacher, lang, indent_mm=6,
        )


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def _setup_page(document):
    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.top_margin = Mm(MARGIN_TOP_MM)
    section.bottom_margin = Mm(MARGIN_BOTTOM_MM)
    section.left_margin = Mm(MARGIN_LEFT_MM)
    section.right_margin = Mm(MARGIN_RIGHT_MM)


def _setup_styles(document):
    style = document.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)


def _render_header(document, meta, lang):
    school = meta.get("school") or ""
    date = meta.get("date") or ""

    meta_p = document.add_paragraph()
    meta_p.add_run(f'{t(lang, "school_class")} ')
    _add_underline_blank(meta_p, school, min_chars=14, size=Pt(11))
    meta_p.add_run(f'   {t(lang, "date")} ')
    _add_underline_blank(meta_p, date, min_chars=10, size=Pt(11))
    meta_p.add_run(f'   {t(lang, "full_name")} ')
    _add_underline_blank(meta_p, "", min_chars=28, size=Pt(11))
    for run in meta_p.runs:
        if run.font.size is None:
            run.font.size = Pt(11)
    meta_p.paragraph_format.space_after = Pt(6)

    title_p = document.add_paragraph()
    run = title_p.add_run(meta.get("title") or t(lang, "default_title"))
    run.font.bold = True
    run.font.size = Pt(16)
    title_p.paragraph_format.space_after = Pt(0)

    subtitle_p = document.add_paragraph()
    run = subtitle_p.add_run(f'{meta.get("subject", "")} · {meta.get("grade", "")} {t(lang, "grade_suffix")}')
    run.font.size = Pt(10.5)
    subtitle_p.paragraph_format.space_after = Pt(6)

    last_header_p = subtitle_p
    if meta.get("instructions"):
        instr_p = document.add_paragraph()
        run = instr_p.add_run(meta["instructions"])
        run.font.italic = True
        instr_p.paragraph_format.space_after = Pt(6)
        last_header_p = instr_p

    _set_paragraph_border(last_header_p, ["bottom"], size=10, space=4)


def build_document_docx(meta, tasks, is_teacher):
    lang = meta.get("language", "ru")
    document = Document()
    _setup_page(document)
    _setup_styles(document)
    _render_header(document, meta, lang)
    for i, task in enumerate(tasks, start=1):
        render_task_docx(document, i, task, is_teacher, lang)

    footer_p = document.sections[0].footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(t(lang, "footer_text"))
    run.font.size = Pt(9)
    return document


def build_docx_documents(meta, tasks):
    return build_document_docx(meta, tasks, is_teacher=False), build_document_docx(meta, tasks, is_teacher=True)


def main():
    import argparse
    import sys

    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("workspace", help="Path to the worksheet draft folder (contains meta.json, tasks/)")
    args = parser.parse_args()

    meta, tasks = load_workspace(args.workspace)
    student_doc, teacher_doc = build_docx_documents(meta, tasks)

    output_dir = Path(args.workspace) / "output"
    output_dir.mkdir(exist_ok=True)
    student_doc.save(output_dir / "worksheet-student.docx")
    teacher_doc.save(output_dir / "worksheet-teacher.docx")
    print(f"Wrote {output_dir / 'worksheet-student.docx'}")
    print(f"Wrote {output_dir / 'worksheet-teacher.docx'}")


if __name__ == "__main__":
    main()
